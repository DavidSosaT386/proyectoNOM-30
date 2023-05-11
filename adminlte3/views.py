from django.shortcuts import redirect, render

from adminlte3.models import anexos, encuesta, respuesta,porfile
from .forms import CustomUserCreationForm,anexosForm,perfilForm
from django.contrib.auth import authenticate, login
from django.contrib.auth import update_session_auth_hash
from django.contrib.auth.forms import PasswordChangeForm
from django.contrib import messages
from django.urls import resolve
import sys

# Para la generación y descarga del fichero
import io
# Para utilizar algunas de las funciones de la librería
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from django.http import HttpResponse
from datetime import datetime
import os
from docx.oxml.shared import OxmlElement, qn
from docx.oxml import ns
from docx.shared import Pt,RGBColor
from docx.oxml.ns import qn
from docx import Document
from docx.enum.style import (
    WD_BUILTIN_STYLE, WD_STYLE, WD_STYLE_TYPE
)
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.dml import MSO_COLOR_TYPE
import docx


global incluye

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def conversorOne(banderola2, banderola4, banderola5):
    if banderola2 == "clas1":
        banderola2 = "Localización"

        banderola4 = conversorC1(banderola4,banderola5)


    if banderola2 == "clas2":
        banderola2 = "Condiciones de la instalación"

        banderola4 = conversorC2(banderola4,banderola5)


    if banderola2 == "clas3":
        banderola2 = "Riesgos inherentes a la instalación"

        banderola4 = conversorC3(banderola4,banderola5)

    if banderola2 == "clas4":
        banderola2 = "Protección contra incendio"

        banderola4 = conversorC4(banderola4,banderola5)

    if banderola2 == "clas5":
        banderola2 = "Peligros circundantes al centro de trabajo"

        banderola4 = conversorC5(banderola4,banderola5)

    if banderola2 == "clas6":
        banderola2 = "Procedimientos administrativos"

        banderola4 = conversorC6(banderola4,banderola5)


    if banderola2 == "clas7":
        banderola2 = "Requerimientos normativos en materia de seguridad"

        banderola4 = conversorC7(banderola4,banderola5)

    return banderola2,banderola4

def conversorC1(banderola4, banderola5):

    if banderola4 == "clas1":
        banderola4 = "Terremotos"

    if banderola4 == "clas2":
        banderola4 = "Clima / Riesgos naturales"

    if banderola4 == "clas3":
        banderola4 = "Riesgos de terceros"

    return banderola4

def conversorC2(banderola4, banderola5):

    if banderola4 == "clas4":
        banderola4 = "Distribución"

    if banderola4 == "clas5":
        banderola4 = "Antiguedad / Condiciones físicas"

    if banderola4 == "clas6":
        banderola4 = "Orden y Limpieza"

    if banderola4 == "clas7":
        banderola4 = "Señalización"

    if banderola4 == "clas8":
        banderola4 = "Seguridad y vigilancia"

    if banderola4 == "clas24":
        banderola4 = "Drenaje"

    return banderola4

def conversorC3(banderola4, banderola5):

    if banderola4 == "clas9":
        banderola4 = "Riesgos inherentes"

    if banderola4 == "clas26":
        banderola4 = "Caseta de control de la subestación"

    if banderola4 == "clas27":
        banderola4 = "Tableros y gabinetes de control"

    if banderola4 == "clas28":
        banderola4 = "Tablero de corriente directa y corriente alterna"

    if banderola4 == "clas29":
        banderola4 = "Cargadores de baterías"

    if banderola4 == "clas30":
        banderola4 = "Cuarto de baterías"

    if banderola4 == "clas31":
        banderola4 = "Cuarto de control supervisorio"

    if banderola4 == "clas32":
        banderola4 = "Cuarto de comunicaciones"

    if banderola4 == "clas33":
        banderola4 = "Transformadores, autotransformadores y reactores"

    if banderola4 == "clas34":
        banderola4 = "Interruptores"

    if banderola4 == "clas35":
        banderola4 = "Transformadores de instrumentos (TC´s, TP´s y DP´s)"

    if banderola4 == "clas36":
        banderola4 = "Cuchillas, apartarrayos, buses y banco de capacitores"

    if banderola4 == "clas37":
        banderola4 = "Trincheras y ductos de cables"

    if banderola4 == "clas38":
        banderola4 = "Subestación de servicios propios"

    if banderola4 == "clas39":
        banderola4 = "Predio de la subestación"

    return banderola4

def conversorC4(banderola4, banderola5):

    if banderola4 == "clas10":
        banderola4 = "Prevención del fuego"

    if banderola4 == "clas11":
        banderola4 = "Reducción del crecimiento y propagación del fuego"

    if banderola4 == "clas12":
        banderola4 = "Sistema de detección y alarma de incendios"

    if banderola4 == "clas13":
        banderola4 = "Sistemas fijos de extinción"

    if banderola4 == "clas14":
        banderola4 = "Extintores portátiles y móviles"

    if banderola4 == "clas15":
        banderola4 = "Brigada contra incendio"

    if banderola4 == "clas16":
        banderola4 = "Convenios de cooperación y ayuda mutua"

    if banderola4 == "clas40":
        banderola4 = "Sistema de detección y alarma de incendio en caseta de control"

    if banderola4 == "clas41":
        banderola4 = "Sistema fijo contraincendio para transformadores y autotransformadores"

    if banderola4 == "clas42":
        banderola4 = "Agentes pasivos"

    return banderola4

def conversorC5(banderola4, banderola5):

    if banderola4 == "clas17":
        banderola4 = "Peligros circundantes"

    return banderola4

def conversorC6(banderola4, banderola5):

    if banderola4 == "clas18":
        banderola4 = "Procedimientos de mantenimiento e inspección"

    if banderola4 == "clas19":
        banderola4 = "Procedimientos de seguridad"

    if banderola4 == "clas20":
        banderola4 = "Protección civil / planes de emergencia / y se cuenta con los procedimientos por escrito para"

    if banderola4 == "clas21":
        banderola4 = "Seguridad ambiental"

    if banderola4 == "clas22":
        banderola4 = "Capacitacion y entrenamiento"

    if banderola4 == "clas23":
        banderola4 = "Procedimientos para control de contratistas"

    return banderola4

def conversorC7(banderola4, banderola5):

    if banderola4 == "clas15":
        banderola4 = "Brigada contra incendio"

    if banderola4 == "clas16":
        banderola4 = "Convenios de cooperación y ayuda mutua"

    if banderola4 == "clas19":
        banderola4 = "Procedimientos de seguridad"

    if banderola4 == "clas20":
        banderola4 = "Protección civil / planes de emergencia / y se cuenta con los procedimientos por escrito para"

    if banderola4 == "clas21":
        banderola4 = "Seguridad ambiental"

    if banderola4 == "clas22":
        banderola4 = "Capacitacion y entrenamiento"

    if banderola4 == "clas23":
        banderola4 = "Procedimientos para control de contratistas"

    if banderola4 == "clas43":
        banderola4 = "Programa de trabajo de seguridad, salud y prevención de riesgos"

    return banderola4
    

def preguntas(usado,document):
    
    banderola3 = " "
    banderola5 = " "
    preg = respuesta.objects.filter(usuario = usado).values('pregunta', 'condicion', 'observacion')

    numerolo = 0

    
    for item in preg:

        numerolo = numerolo + 1
        
        num = item['condicion']
        rpreg = encuesta.objects.filter(id = item['pregunta']).values('pregunta').first()
        rpreg2 = encuesta.objects.filter(id = item['pregunta']).values('clasificacion','sub_clasificacion').first()
        rpreg3 = encuesta.objects.filter(id = item['pregunta']).values('clasificacion','sub_clasificacion').first()
        #instancia = {}
        #rpreg3.get_clasificacion_sudisplay()

        banderola2 = rpreg3.get('clasificacion')
        banderola4 = rpreg3.get('sub_clasificacion')
        

        banderola = rpreg.get('pregunta')


        container = encuesta.objects.filter()

        banderola2,banderola4 = conversorOne(banderola2,banderola4, banderola5)
        #banderola4 = conversorTwo(banderola4)


        if banderola3 != banderola2:
            document.add_heading(banderola2,1)
            banderola3 = banderola2

        if banderola5 != banderola4:
            numerolo = 1
            document.add_heading(banderola4,2)
            banderola5 = banderola4

        numerolo2 = str(numerolo)

        document.add_paragraph(numerolo2 + '.- ' + banderola, style = 'Pregunta')

 

        if item['condicion'] == False:

            
            res = item['observacion']
            document.add_paragraph(res, style = 'Respuesta')

        elif item['condicion'] == True:

            document.add_paragraph("No hay condiciones subestandares", style = 'Respuesta')

    paranexo = respuesta.objects.filter(usuario = usado, no_reporte = 1).values('no_reporte').first()

    infonexo = anexos.objects.filter(no_reporte = paranexo['no_reporte'], usuario = usado).values('foto_anexo','descripcion')

    document.add_page_break() 

    

    document.add_heading("XI.- Conclusiones", 1)
    document.add_heading("XII.- Reporte Fotografico", 1)


    for x in infonexo:
        
        path = os.path.realpath('media/' + x['foto_anexo'])
        document.add_picture(path, width=Cm(13.77), height=Cm(10.33))
        document.add_paragraph(x['descripcion'],style = 'Fotos')   




def estilos(styles):
    
    mi_p_estilo = styles.add_style('Estilo Parrafo', WD_STYLE_TYPE.PARAGRAPH)
    
    mi_p_estilo.base_style = styles['Normal']
    mi_p_estilo.paragraph_format.space_after = Pt(20)
    mi_p_estilo.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    mi_p_estilo.font.color.rgb = RGBColor(4, 158, 86)
    mi_p_estilo.font.bold = True
    mi_p_estilo.font.size = Pt(18)
    mi_p_estilo.font.name = "Verdana"


    mi_p_estilo = styles.add_style('Respuesta', WD_STYLE_TYPE.PARAGRAPH)
    
    mi_p_estilo.base_style = styles['Estilo Parrafo']
    mi_p_estilo.paragraph_format.space_after = Pt(10)
    mi_p_estilo.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    mi_p_estilo.font.color.rgb = RGBColor(0, 0, 0)
    mi_p_estilo.font.bold = False
    mi_p_estilo.font.size = Pt(12)

    mi_p_estilo = styles.add_style('Fotos', WD_STYLE_TYPE.PARAGRAPH)
    
    mi_p_estilo.base_style = styles['Respuesta']
    mi_p_estilo.paragraph_format.space_after = Pt(10)
    mi_p_estilo.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    mi_p_estilo.font.color.rgb = RGBColor(0, 0, 0)
    mi_p_estilo.font.italic = True
    mi_p_estilo.font.size = Pt(9)

    mi_p_estilo = styles.add_style('general', WD_STYLE_TYPE.PARAGRAPH)
    
    mi_p_estilo.base_style = styles['Respuesta']
    mi_p_estilo.paragraph_format.space_before = Pt(10)
    

    mi_p_estilo = styles.add_style('Pregunta', WD_STYLE_TYPE.PARAGRAPH)
    
    mi_p_estilo.base_style = styles['Estilo Parrafo']
    mi_p_estilo.paragraph_format.space_after = Pt(0)
    mi_p_estilo.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    mi_p_estilo.font.color.rgb = RGBColor(0, 0, 0)
    mi_p_estilo.font.size = Pt(12)


    mi_p_estilo2 = styles.add_style('Estilo Parrafo2', WD_STYLE_TYPE.PARAGRAPH)
    
    mi_p_estilo2.base_style = styles['Estilo Parrafo']
    mi_p_estilo2.font.size = Pt(12)

    mi_p_estilo3 = styles.add_style('Estilo Parrafo3', WD_STYLE_TYPE.PARAGRAPH)
    
    mi_p_estilo3.base_style = styles['Estilo Parrafo']
    mi_p_estilo3.font.color.rgb = RGBColor(0, 0, 255)
    mi_p_estilo3.font.size = Pt(14)
    mi_p_estilo3.font.name = "Arial"



    mi_p_estilo3 = styles.add_style('Encabezado', WD_STYLE_TYPE.PARAGRAPH)
    
    mi_p_estilo3.base_style = styles['Normal']
    mi_p_estilo3.paragraph_format.space_after = Pt(0)
    mi_p_estilo3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    mi_p_estilo3.font.color.rgb = RGBColor(4, 158, 86)
    mi_p_estilo3.font.bold = True
    mi_p_estilo3.font.size = Pt(10)
    mi_p_estilo3.font.name = "Helvetica LT Std Cond"



    mi_p_estilo4 = styles.add_style('Encabezado2', WD_STYLE_TYPE.PARAGRAPH)
    
    mi_p_estilo4.base_style = styles['Encabezado']
    mi_p_estilo4.font.bold = False
    mi_p_estilo4.font.size = Pt(9)

def generalidades(document,centro):

    document.add_heading("I.- Introducción.", 1)
    document.add_paragraph('        Como parte del programa de trabajo de la G.R.T Sureste, AGREGAR FECHA,'
    ' se efectuó una visita de evaluación de riesgos a la Zona de Transmisión ' + centro +
    '\n    El presente informe contiene los resultados de la verificación de la gestión y del estado que guardan las '
    'diversas áreas de la instalación. Así mismo, contiene las observaciones y recomendaciones necesarias para eliminar '
    'o reducir y controlar las condiciones de riesgo existentes.', style = 'general')

    
    document.add_heading("II.- Objetivo.", 1)
    document.add_paragraph('        Determinar las condiciones subestándar prevalecientes en los sistemas,'
    ' equipos, procedimientos y controles, que puedan ser un factor de riesgo en la subestación,'
    ' y apoyar al personal de la Zona de Transmisión ' + centro +
    ' en la determinación de acciones alternativas para su solución.', style = 'general')

    document.add_heading("III.- Normatividad aplicada.", 1)
    document.add_paragraph('NOM-030-STPS-2009           Servicios preventivos de seguridad y salud en el trabajo.'
    'RFSST                       Reglamento Federal de Seguridad y Salud en el Trabajo', style = 'general')
    
    document.add_heading("IV.-  Antecedentes.", 1)
    document.add_paragraph('        Con base en la NOM-030-STPS-2009 Servicios preventivos de seguridad y salud en el trabajo,' 
        '"Guía de Verificación para identificación de áreas de oportunidad", y el "Programa de Actividades de Gerencia ' 
        'Regional de Transmisión Sureste, Departamento de Control de Gestión,--- AGREGAR FECHA ---, se efectuó una visita ' 
        'para efectuar un diagnóstico integral sobre las condiciones de seguridad y salud en el trabajo de la Zona de Transmisión ' + centro + 
        '\n    Los conceptos abordados incluyen la verificación de la gestión y el estado que guarda la instalación.', style = 'general')


    document.add_heading("V.- Información General de la Instalación.", 1)

    document.add_heading("VI.-  Mejoras realizadas en materia de Seguridad y Salud en el Trabajo.", 1)

    

import string

def calcMes(mes):

    if mes == 'January':
        mes2 = 'Enero'
    if mes == 'February':
        mes2 = 'Febrero'
    if mes == 'March':
        mes2 = 'Marzo'
    if mes == 'April':
        mes2 = 'Abril'
    if mes == 'July':
        mes2 = 'Julio'
    if mes == 'October':
        mes2 = 'Octubre'
    if mes == 'December':
        mes2 = 'Diciembre'
    if mes == 'November':
        mes2 = 'Noviembre'
    if mes == 'August':
        mes2 = 'Agosto'
    if mes == 'September':
        mes2 = 'Septiembre'
    if mes == 'June':
        mes2 = 'Junio'
    if mes == 'May':
        mes2 = 'Mayo'

    return mes2

def coverletter_export(request,establis,no_reporte):
    
    establecimiento = establis
    probando = "QWERTYUOPASDFGHJKLÑZXCVBNM"
    abe = "qwertyuiopasdfghjklñzxcvbnm"
    establedo = establis + " "
    k = 0
    index = " "

    for letra in establis:
        k = k + 1
        
        if letra in probando:
            index = establecimiento.find(letra)
            establecimiento = establecimiento[:index] + ' ' + establecimiento[index:]

        #if (letra == 'I' and establis[k-1] in abe) or (letra == 'I' and establis[k+1] == 'I'):

        if letra == 'I' and establedo[k] in abe or letra == 'I' and establedo[k] == 'I':
            index = establecimiento.find(letra)
            establecimiento = establecimiento[:index] + ' ' + establecimiento[index:]


        



    n_report = no_reporte
    us = porfile.objects.filter(usuario= request.user ).first()  
    if request.method == "POST":
        user1 = request.user
        
        ape = porfile.objects.filter(usuario = user1 ).values('nombre','apellidos','centro_trabajo').first() 
        nom = ape['nombre'] + " " + ape['apellidos']
        centro = ape['centro_trabajo']

        print(nom) 
        pata = os.path.realpath('adminlte3/prueba.docx')
        document = Document()
        styles = document.styles
        estilos(styles)

        add_page_number(document.sections[0].footer.paragraphs[0].add_run())

        seccion = document.sections[0].header
        encabezado = seccion.paragraphs[0]

        
        seccion.add_paragraph("CFE Transmisión", style = "Encabezado")
        seccion.add_paragraph("Gerencia Regional de Transmisión Sureste \nDepartamento de Control de Gestión", style = "Encabezado2")

        

        cfe = encabezado.add_run() 

        path = os.path.realpath('adminlte3/static/ImPersonal/CFE.jpg')

        cfe.add_picture(path, width=Cm(7.01), height=Cm(1.49))

        

        t1 = "INFORME"
        t2 = "NOM-030-STPS-2009\n DIAGNOSTICO INTEGRAL SOBRE CONDICIONES DE SEGURIDAD \n EN LA ZONA DE TRANSMISION " + centro.upper()
        t3 = "INCLUYE"

        mes = datetime.today().strftime('%B')
        
        fecha = datetime.today().strftime('%d de ' + calcMes(mes) +' del %Y')

        



        document.add_paragraph("\n\n\n\n")

        establecimiento = establecimiento[0].upper() + establecimiento[1:]




        document.add_paragraph(t1, style = "Estilo Parrafo")
        document.add_paragraph(t2, style = "Estilo Parrafo2")
        document.add_paragraph(t3, style = "Estilo Parrafo3")
        document.add_paragraph(establecimiento, style = "Estilo Parrafo3")
        document.add_paragraph("ING. " + nom, style = "Estilo Parrafo2")

        document.add_paragraph(fecha)

        #--------------------------------------salto de pagina--------------------------------------------------------------------

        document.add_page_break()

        generalidades(document,centro)  

        #impresion de preguntas
        preguntas(user1,document)

        

        



        document_data = io.BytesIO()
        document.save(document_data)
        document_data.seek(0)
        response = HttpResponse( document_data.getvalue(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",)
        response["Content-Disposition"] = 'attachment; filename = "INFORME DIAGNOSTICO INTEGRAL NOM-030.docx"'
        response["Content-Encoding"] = "UTF-8" 
        return response
        
    else:
        return render(request,'adminlte/gestionNOM30/generarReporte.html',{'us':us, 'establecimiento':establecimiento, 'n_reporte':n_report}  )    





import os
from django.conf import settings

def home(request):
    
    
    user1 = request.user
    if request.user.is_anonymous:
        us="hola"
    else:
        n_report= 1
        respuestas =  respuesta.objects.filter(usuario=request.user)
        if respuestas:
            respuestas.delete()      
  
        anex = anexos.objects.filter(no_reporte=n_report, usuario= request.user)
        if anex:
            for a in anex:
                an= anexos.objects.get(id=a.id)
                an.delete()
        us = porfile.objects.filter(usuario= user1 ).first()   

    return render(request,'adminlte/index.html',{'us':us} )

def registro(request):
    data = {
        'form':CustomUserCreationForm()

    }
    if request.method == 'POST':
        formulario = CustomUserCreationForm(data=request.POST)
        if formulario.is_valid():
            formulario.save()
            user=authenticate(username= formulario.cleaned_data['username'],password= formulario.cleaned_data['password1'])
            login(request, user)
            nombre=request.POST.get("nombre")
            apellidos=request.POST.get("apellidos")
            foto_perfil=request.FILES.get("foto")
            centro_trabajo=request.POST.get("centro")
            user1 = request.user
            g= porfile(usuario=user1, nombre=nombre, apellidos=apellidos, foto_perfil=foto_perfil, centro_trabajo=centro_trabajo )
            g.save()
            messages.success(request,"Registro correcto")
            return redirect(to="index")
        data["form"] =  formulario     
    return render(request,'registration/registro.html',data )    



def almacen(request, alma):
    almac =  alma
    us = respuesta.objects.filter(usuario = request.user).last()
    usi = porfile.objects.filter(usuario=request.user).first()
    if us == None:
        n_reporte = 1
    else:    
        n_reporte = us.no_reporte + 1
    quis = encuesta.objects.filter(categoria='categoria1')
    if request.method == 'POST':
        user = request.user
        a= zip(
            list(quis),
            request.POST.getlist("condiciones"),
            request.POST.getlist("observaciones"),

            )
        dicti= [{'pregu': qu, 'condi': condiciones , 'obser': observaciones } for qu, condiciones, observaciones in a ]
        for dd in dicti:
            g= respuesta(no_reporte=n_reporte,usuario= user, pregunta=dd['pregu'], condicion= dd['condi'], observacion= dd['obser'])
            g.save()
        return redirect(Anexos, n_reporte,almac )
    else:        
        contexto={
            'pregu': quis,
            'us':usi,
            'alma':almac,
        }
    return render(request, 'adminlte/gestionNOM30/almacen.html', contexto)

def oficina(request, off):

    office =  off
    us = respuesta.objects.filter(usuario = request.user).last()
    usi = porfile.objects.filter(usuario=request.user).first()
    if us == None:
        n_reporte = 1
    else:    
        n_reporte = us.no_reporte + 1
    quis = encuesta.objects.filter(categoria='categoria2')
    if request.method == 'POST':
        user = request.user
        a= zip(
            list(quis),
            request.POST.getlist("condiciones"),
            request.POST.getlist("observaciones"),

            )
        dicti= [{'pregu': qu, 'condi': condiciones , 'obser': observaciones } for qu, condiciones, observaciones in a ]
        for dd in dicti:
            g= respuesta(no_reporte=n_reporte,usuario= user, pregunta=dd['pregu'], condicion= dd['condi'], observacion= dd['obser'])
            g.save()
        return redirect(Anexos, n_reporte, office )
    else:        
        contexto={
            'pregu': quis,
            'us':usi,
            'ofis':office,
        }
    return render(request, 'adminlte/gestionNOM30/oficina.html', contexto)    

def subestacion(request, sub):
    subesta =  sub
    us = respuesta.objects.filter(usuario = request.user).last()
    usi = porfile.objects.filter(usuario=request.user).first()
    if us == None:
        n_reporte = 1
    else:    
        n_reporte = us.no_reporte + 1
    quis = encuesta.objects.filter(categoria='categoria3')
    if request.method == 'POST':
        user = request.user
        a= zip(
            list(quis),
            request.POST.getlist("condiciones"),
            request.POST.getlist("observaciones"),

            )
        dicti= [{'pregu': qu, 'condi': condiciones , 'obser': observaciones } for qu, condiciones, observaciones in a ]
        for dd in dicti:
            g= respuesta(no_reporte=n_reporte,usuario= user, pregunta=dd['pregu'], condicion= dd['condi'], observacion= dd['obser'])
            g.save()
        return redirect(Anexos, n_reporte, subesta )
    else:        
        contexto={
            'pregu': quis,
             'us':usi,
             'sub':subesta,
        }
    return render(request, 'adminlte/gestionNOM30/subestacion.html', contexto) 


def Anexos(request, n_reporte, establis):

    
    form = anexosForm()
    usi = porfile.objects.filter(usuario=request.user).first()
    contexto={
            'form':form,
            'us':usi,
        }
    if request.method == 'POST':
        
        no_reporte= n_reporte
        a=zip(
             request.FILES.getlist('foto_anexo'),
             request.POST.getlist('descripcion')
            
        )
        dicti= [{'foto': foto_anexo, 'descrip': descripcion} for foto_anexo, descripcion in a ]
        for dd in  dicti:
            g= anexos(usuario=request.user, no_reporte = no_reporte, foto_anexo = dd['foto'], descripcion= dd['descrip'])
            g.save()
        messages.success(request, 'Reporte finalizado') 

      
        return redirect(coverletter_export,establis,n_reporte)    
    return render(request, 'adminlte/gestionNOM30/anexos.html', contexto)

def change(request):
    usi = porfile.objects.filter(usuario=request.user).first()
    if request.method == 'POST':
        form = PasswordChangeForm(request.user, request.POST)
        if form.is_valid():
            user = form.save()
            update_session_auth_hash(request, user) 
            messages.success(request, 'Su contraseña ha sido actualizada')
            contexto ={
                'form': form, 'us':usi
            }
            return redirect(to="index")
        else:
            messages.error(request, 'Verifique y corrija los datos ingresados')
            tipo= "error"
            titulo= "Error"
            contexto = {
                'form': form, 'us':usi , 'tipo':tipo, 'titulo': titulo
            }



    else:
        usi = porfile.objects.filter(usuario=request.user).first()
        form = PasswordChangeForm(request.user)
        contexto ={
                'form': form, 'us':usi
            }
    return render(request, 'registration/change.html', contexto)




def cambiarPerfil(request):
    usi = porfile.objects.filter(usuario=request.user).first()
    employed= porfile.objects.get(usuario = request.user)
    if request.method == 'POST':
        form = perfilForm(request.POST, request.FILES, instance= employed)
        contexto={
            'form':form
        }   
        if form.is_valid:
            form.save()
            messages.success(request, "El usuario ha sido editado con exito")
            return redirect(to="index")
    else:
        form =perfilForm(instance= employed)
        contexto ={
            'form':form,
            'us':usi,
        }
    return render(request, 'registration/changeprofile.html', contexto)    

